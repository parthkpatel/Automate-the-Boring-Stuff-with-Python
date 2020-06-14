# A program that, given a .txt file path, goes through the text file line by line (will assume one guest per line)
# and appends a page to a document containing a custom invitation for that guest.
# The program will then save this document of invitations in the same folder as the given .txt file

import pyinputplus as pyip, docx, os
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date


def custom_invitations(path):
    document = docx.Document()

    # Adding brush script style
    brush_script_style = document.styles.add_style('Brush Script Style', WD_STYLE_TYPE.CHARACTER)
    brush_script_style.font.size = Pt(20)
    brush_script_style.font.name = 'Brush Script MT'

    # Adding Guest style
    guest_style = document.styles.add_style('Guest', WD_STYLE_TYPE.CHARACTER)
    guest_style.font.size = Pt(25)
    guest_style.font.bold = True

    # Adding Date Style
    date_style = document.styles.add_style('Date', WD_STYLE_TYPE.CHARACTER)
    date_style.font.size = Pt(20)

    guest_list_file = open(path)
    guest_list_text = guest_list_file.read()
    guest_list = guest_list_text.split('\n')

    # Quick check to make sure that there is at least one guest in the list
    if len(guest_list) == 0:
        return

    for i, guest in enumerate(guest_list):
        invitation_paragraph = document.add_paragraph()
        invitation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        invitation_paragraph.add_run('It would be a pleasure to have the company of', style='Brush Script Style')

        guest_paragraph = document.add_paragraph()
        guest_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        guest_paragraph.add_run(guest, style='Guest')

        location_paragraph = document.add_paragraph()
        location_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        location_paragraph.add_run('at 11010 Memory Lane on the Evening of', style='Brush Script Style')

        date_paragraph = document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_paragraph.add_run(date.today().strftime("%B %d"), style='Date')

        time_paragraph = document.add_paragraph()
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_paragraph.add_run("at 7 o'clock", style='Brush Script Style')

        # Do not want a page break after last guest
        if i != len(guest_list) - 1:
            document.add_page_break()

    document.save(os.path.join(os.path.dirname(path), 'invites.docx'))
    guest_list_file.close()


if __name__ == '__main__':
    path_prompt = 'Please enter the full path to the guest list .txt file:\n'
    input_path = pyip.inputFilepath(prompt=path_prompt)
    custom_invitations(input_path)
